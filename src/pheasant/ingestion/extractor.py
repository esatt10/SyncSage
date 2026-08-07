"""Document text extraction for PDF/DOCX/PPTX/XLSX/DOC/RTF/EPUB/HTML ingestion.

This module owns PDF, DOCX and HTML and the provider/config surface for all
formats; :mod:`pheasant.ingestion.office` adds PPTX, XLSX, EPUB and RTF, and
:mod:`pheasant.ingestion.msdoc` the legacy binary DOC. ``extract_builtin`` is
the one place formats map to readers.

A ``.pdf`` or ``.docx`` file carries plenty of indexable text, but none of it
is reachable with :meth:`pathlib.Path.read_text`. Before this module the
ingestion pipeline **accepted** both extensions (``parse_file`` /
``parse_connector_payload`` admit them, ``content_types.artifact_type``
labels them ``"document"``) and then threw the content away:
``read_text`` / ``read_text_bytes`` hard-returned ``""`` for both, so a PDF
was discovered, hashed, given an artifact row and a graph node — and produced
**zero chunks**. It was findable by path and invisible by content. The two
dependencies needed to fix it (``pymupdf``, ``python-docx``) were already
declared as *core* requirements in ``pyproject.toml`` and imported nowhere.

This module closes that gap the same way Synapse 25.4 closed it for images
(:mod:`pheasant.ingestion.captioner`) and audio
(:mod:`pheasant.ingestion.transcriber`): a bytes -> indexable-text provider
with a deterministic offline default, selected by config, built only when a
source actually admits the extension. The shared sidecar/fingerprint helpers
in :mod:`pheasant.ingestion._modal` are reused rather than reimplemented.

Where it deliberately differs from 25.4
---------------------------------------
Captioning and transcription need a *model* to invent text that isn't in the
bytes, so their real providers make network calls. Document extraction needs
no model and no network: the text is already in the file, and pulling it out
is pure computation. That makes this the **stronger** guarantee — every
provider here is fully offline and deterministic, so rule 1 ("never put an
LLM call in the indexing path") holds by construction rather than by
convention. There is no network path to gate.

Providers
---------
- :class:`BuiltinExtractor` — **pure standard library, zero third-party
  dependencies.** PDF: locate content streams, ``zlib``-inflate the
  FlateDecode ones, and scan the PostScript-like operator stream for
  text-showing operators. DOCX: it is a ZIP, so ``zipfile`` +
  ``xml.etree`` over ``word/document.xml``. HTML: BeautifulSoup when
  present (already a core dep), else a regex strip. This exists because
  the repo ships a dependency-light shim environment (root ``yaml.py``,
  the ``pheasant/`` path bridge) in which ``pymupdf`` may be absent — and
  because a text extractor that dies when an optional import is missing
  would take the whole sync down with it.
- :class:`NativeExtractor` — ``pymupdf`` for PDF and ``python-docx`` for
  DOCX (both already core deps), with higher fidelity than the builtin on
  hard PDFs: CID/Type0 fonts, custom encodings, complex layouts. Falls
  back to the builtin per-format when its import is unavailable.
- :class:`SandboxedExtractor` — the builtin PDF *tokenizer* run inside the
  Phase-34 WASM sandbox under a fuel + linear-memory cap with **zero host
  capabilities**. See :mod:`pheasant.ingestion.extractor_sandbox`; the
  threat model is documented there.
- :class:`AutoExtractor` — the default. Native when importable, builtin
  otherwise, and ``""`` if both somehow fail. Never raises into the sync.

Priority order for every format, mirroring 25.4 exactly: an authored
``<file>.extract.txt`` sidecar wins over any provider, so a fixture or an
air-gapped deployment can supply exact text with no extractor at all.

Idempotency note
----------------
Extraction happens inside :func:`pheasant.ingestion.pipeline.parse_file` and
``parse_connector_payload``, but the sync engine's pre-read skip
(``_can_skip_before_read``) compares the file's content sha256 *before*
reading bytes, so an unchanged PDF in an incremental sync is skipped and
**never re-extracted** — the same zero-work guarantee the embedder (21.4),
captioner (25.4A) and transcriber (25.4B) get. Extraction is a pure function
of the bytes, so a re-sync of unchanged content reproduces identical chunks.
"""

from __future__ import annotations

import logging
import re
import zlib
from pathlib import Path
from typing import TYPE_CHECKING, Any, Protocol, runtime_checkable

from pheasant.ingestion._modal import sidecar_text as _sidecar_text

if TYPE_CHECKING:
    from pheasant.config.schema import ExtractorSettings

logger = logging.getLogger(__name__)

EXTRACT_SIDECAR_SUFFIX = ".extract.txt"

PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
PPTX_EXTENSIONS = {".pptx"}
XLSX_EXTENSIONS = {".xlsx"}
DOC_EXTENSIONS = {".doc"}
RTF_EXTENSIONS = {".rtf"}
EPUB_EXTENSIONS = {".epub"}
HTML_EXTENSIONS = {".html", ".htm", ".xhtml"}

# Every extension this module can turn into text, HTML aside (HTML is opt-in;
# see `extractable_extensions`). Kept as one set so `source_includes_documents`
# and `content_types.DOCUMENT_EXTENSIONS` cannot drift apart: if a format is
# accepted by the pipeline but missing from here, no extractor gets built for a
# source that only carries that format, and the file indexes with zero
# chunks — exactly the failure this module exists to remove.
EXTRACTED_EXTENSIONS = (
    PDF_EXTENSIONS
    | DOCX_EXTENSIONS
    | PPTX_EXTENSIONS
    | XLSX_EXTENSIONS
    | DOC_EXTENSIONS
    | RTF_EXTENSIONS
    | EPUB_EXTENSIONS
)

# Guard against a pathological input exhausting host memory. A single
# inflated content stream is capped, and so is the total extracted text: a
# "zip bomb" PDF whose streams inflate to gigabytes is bounded here rather
# than in the OOM killer. Both are generous for real documents.
MAX_STREAM_INFLATED_BYTES = 32 * 1024 * 1024
MAX_TOTAL_TEXT_CHARS = 8 * 1024 * 1024

_DOCX_MAIN_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


@runtime_checkable
class DocumentExtractor(Protocol):
    """Document-bytes -> indexable-text provider. ``calls`` counts extractions."""

    model: str
    calls: int

    def extract(self, content: bytes, relative_path: str) -> str: ...


# ---------------------------------------------------------------------------
# PDF: pure-stdlib content-stream text extraction
# ---------------------------------------------------------------------------
#
# A PDF's visible text lives in *content streams* — PostScript-like operator
# programs — which are usually compressed with FlateDecode (i.e. zlib, which
# is in the standard library). Extraction is therefore two independent steps,
# kept as separate functions so each is testable in isolation and so the
# tokenizer can be ported verbatim to the WASM guest:
#
#   1. `iter_pdf_content_streams` — carve the streams out of the file and
#      inflate them.
#   2. `scan_pdf_content_stream`  — walk one inflated operator program and
#      pull out the text-showing operators' string operands.
#
# Fidelity limits, stated honestly: this handles uncompressed and
# FlateDecode streams with simple (single-byte) font encodings, which covers
# the overwhelming majority of real-world text PDFs. It does NOT decrypt
# encrypted PDFs, decode LZW/CCITT/JBIG2 streams, or resolve Type0/CID font
# CMaps (whose multi-byte codes need the font's embedded mapping table). For
# those, `NativeExtractor` (pymupdf) is materially better — which is exactly
# why `AutoExtractor` prefers it and why this is the fallback, not the
# default. A PDF this cannot read yields less text, never a crash.

_STREAM_KEYWORD = re.compile(rb"(?<![A-Za-z])stream[\r\n]")
_LENGTH_RE = re.compile(rb"/Length\s+(\d+)")


def _inflate(data: bytes) -> bytes | None:
    """zlib-inflate ``data``, tolerating trailing garbage.

    Content-stream slicing is approximate (see ``iter_pdf_content_streams``),
    so the buffer handed here often runs past the deflate stream's true end.
    ``decompressobj`` stops cleanly at that boundary and parks the remainder
    in ``unused_data`` instead of raising, which is what makes the cheap
    slicing safe. Output is capped to bound a decompression bomb.
    """
    try:
        obj = zlib.decompressobj()
        out = obj.decompress(data, MAX_STREAM_INFLATED_BYTES)
        return out or None
    except zlib.error:
        # Not a deflate stream (raw/LZW/DCT image data, or a mis-sliced
        # buffer). Not an error worth surfacing — most streams in a PDF are
        # images and fonts, and skipping them is correct.
        return None


def iter_pdf_content_streams(content: bytes) -> list[bytes]:
    """Carve out and inflate every plausibly-textual stream in a PDF.

    Streams are located by the ``stream``/``endstream`` keyword pair rather
    than by following the cross-reference table, so a damaged or
    incrementally-updated xref (common in real files, and the usual reason a
    strict parser gives up entirely) does not prevent extraction. Over-long
    slices are harmless: ``_inflate`` ignores trailing bytes, and an
    uncompressed stream's true end is taken from its ``/Length``.
    """
    streams: list[bytes] = []
    for match in _STREAM_KEYWORD.finditer(content):
        start = match.end()
        end = content.find(b"endstream", start)
        if end == -1:
            end = len(content)
        raw = content[start:end]
        if not raw:
            continue
        inflated = _inflate(raw)
        if inflated is not None:
            streams.append(inflated)
            continue
        # Uncompressed stream: trust /Length from the object dict that
        # precedes the `stream` keyword, falling back to the carved slice.
        header = content[max(0, match.start() - 512) : match.start()]
        declared_lengths = _LENGTH_RE.findall(header)
        if declared_lengths:
            # The last /Length before this `stream` keyword is this object's.
            declared = int(declared_lengths[-1])
            if 0 < declared <= len(raw):
                raw = raw[:declared]
        if b"Tj" in raw or b"TJ" in raw or b"BT" in raw:
            streams.append(raw)
    return streams


def _decode_pdf_literal(raw: bytes) -> bytes:
    r"""Resolve escape sequences inside a PDF ``(...)`` literal string.

    Handles ``\n \r \t \b \f \( \) \\``, three-digit octal codes, and a
    backslash-newline line continuation (which yields nothing).
    """
    out = bytearray()
    i = 0
    simple = {
        ord("n"): 10,
        ord("r"): 13,
        ord("t"): 9,
        ord("b"): 8,
        ord("f"): 12,
        ord("("): 40,
        ord(")"): 41,
        ord("\\"): 92,
    }
    while i < len(raw):
        byte = raw[i]
        if byte != 0x5C:  # not a backslash
            out.append(byte)
            i += 1
            continue
        i += 1
        if i >= len(raw):
            break
        nxt = raw[i]
        if nxt in simple:
            out.append(simple[nxt])
            i += 1
        elif 0x30 <= nxt <= 0x37:  # octal escape, 1-3 digits
            digits = bytearray()
            while i < len(raw) and len(digits) < 3 and 0x30 <= raw[i] <= 0x37:
                digits.append(raw[i])
                i += 1
            out.append(int(digits.decode("ascii"), 8) & 0xFF)
        elif nxt in (10, 13):  # line continuation: emits nothing
            i += 1
            if nxt == 13 and i < len(raw) and raw[i] == 10:
                i += 1
        else:
            out.append(nxt)
            i += 1
    return bytes(out)


def _decode_pdf_text_bytes(raw: bytes) -> str:
    """Decode show-operator string bytes to text.

    A leading UTF-16 BOM marks a wide string (PDF allows either). Otherwise
    the bytes are single-byte codes; ``cp1252`` matches WinAnsiEncoding, the
    default for the standard fonts, and degrades gracefully elsewhere.
    """
    if raw.startswith((b"\xfe\xff", b"\xff\xfe")):
        return raw.decode("utf-16", errors="ignore")
    return raw.decode("cp1252", errors="replace")


def scan_pdf_content_stream(stream: bytes) -> str:
    """Extract shown text from one inflated PDF content stream.

    Walks the operator program tracking only what text extraction needs: the
    most recent string operands, and the operators that show them (``Tj``,
    ``TJ``, ``'``, ``"``) or move the text cursor (``Td``, ``TD``, ``T*``,
    ``ET``) and therefore imply a line break. Everything else — graphics
    state, paths, colour, images — is skipped.

    This is the function the WASM guest mirrors; see
    :mod:`pheasant.ingestion.extractor_sandbox` for the parity guarantee.
    """
    out: list[str] = []
    pending: list[str] = []  # strings collected since the last operator
    i = 0
    n = len(stream)

    def flush_line() -> None:
        if pending:
            out.append("".join(pending))
            pending.clear()
        elif out and out[-1] != "":
            out.append("")

    while i < n:
        byte = stream[i]
        if byte == 0x28:  # '(' literal string
            depth = 1
            i += 1
            start = i
            while i < n and depth:
                if stream[i] == 0x5C:  # backslash escapes the next byte
                    i += 2
                    continue
                if stream[i] == 0x28:
                    depth += 1
                elif stream[i] == 0x29:
                    depth -= 1
                    if depth == 0:
                        break
                i += 1
            literal = stream[start:i]
            i += 1
            pending.append(_decode_pdf_text_bytes(_decode_pdf_literal(literal)))
        elif byte == 0x3C and i + 1 < n and stream[i + 1] != 0x3C:  # '<' hex string
            end = stream.find(b">", i + 1)
            if end == -1:
                break
            hex_digits = bytes(c for c in stream[i + 1 : end] if not chr(c).isspace())
            if len(hex_digits) % 2:
                hex_digits += b"0"
            try:
                pending.append(_decode_pdf_text_bytes(bytes.fromhex(hex_digits.decode("ascii"))))
            except ValueError:
                pass
            i = end + 1
        elif byte == 0x25:  # '%' comment to end of line
            while i < n and stream[i] not in (10, 13):
                i += 1
        elif chr(byte).isalpha() or byte == 0x27 or byte == 0x22:
            start = i
            while i < n and (chr(stream[i]).isalnum() or stream[i] in (0x2A, 0x27, 0x22)):
                i += 1
            op = stream[start:i]
            if op in (b"Tj", b"'", b'"'):
                if pending:
                    out.append("".join(pending))
                    pending.clear()
                if op in (b"'", b'"'):
                    out.append("")
            elif op == b"TJ":
                if pending:
                    out.append("".join(pending))
                    pending.clear()
            elif op in (b"Td", b"TD", b"T*", b"ET"):
                flush_line()
            else:
                pending.clear()
        else:
            i += 1

    if pending:
        out.append("".join(pending))
    return "\n".join(out)


def extract_pdf_text_builtin(content: bytes) -> str:
    """Pure-stdlib PDF text extraction (no third-party dependency)."""
    parts: list[str] = []
    total = 0
    for stream in iter_pdf_content_streams(content):
        text = scan_pdf_content_stream(stream)
        if not text.strip():
            continue
        parts.append(text)
        total += len(text)
        if total >= MAX_TOTAL_TEXT_CHARS:
            logger.warning("PDF extraction hit the %d-char cap; truncating", MAX_TOTAL_TEXT_CHARS)
            break
    return _tidy("\n".join(parts))


# ---------------------------------------------------------------------------
# DOCX: pure-stdlib (a .docx is a ZIP of XML)
# ---------------------------------------------------------------------------


def extract_docx_text_builtin(content: bytes) -> str:
    """Pure-stdlib DOCX text extraction.

    Unlike the PDF path this is not a fidelity compromise: ``word/document.xml``
    stores the body text as ``<w:t>`` runs inside ``<w:p>`` paragraphs, so
    reading it with ``zipfile`` + ``xml.etree`` recovers the document text
    directly. Tables are included (their cells are ordinary paragraphs).
    """
    import io
    import xml.etree.ElementTree as ET
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            names = set(archive.namelist())
            if "word/document.xml" not in names:
                return ""
            xml_bytes = archive.read("word/document.xml")
    except (zipfile.BadZipFile, KeyError, OSError) as exc:
        logger.debug("DOCX archive unreadable: %s", exc)
        return ""

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError as exc:
        logger.debug("DOCX document.xml unparseable: %s", exc)
        return ""

    lines: list[str] = []
    for paragraph in root.iter(f"{_DOCX_MAIN_NS}p"):
        buf: list[str] = []
        for node in paragraph.iter():
            tag = node.tag
            if tag == f"{_DOCX_MAIN_NS}t":
                buf.append(node.text or "")
            elif tag == f"{_DOCX_MAIN_NS}tab":
                buf.append("\t")
            elif tag == f"{_DOCX_MAIN_NS}br":
                buf.append("\n")
        lines.append("".join(buf))
    return _tidy("\n".join(lines))


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


def extract_html_text(content: bytes) -> str:
    """Strip markup so HTML/XML indexes its prose, not its tags.

    ``bs4`` is already a core dependency (the Confluence connector uses it),
    but a regex strip stands in when it is missing so the dependency-light
    environment still degrades to something useful.
    """
    raw = content.decode("utf-8", errors="ignore")
    try:
        from bs4 import BeautifulSoup
    except ModuleNotFoundError:
        no_script = re.sub(r"<script[\s\S]*?</script>", " ", raw, flags=re.I)
        no_style = re.sub(r"<style[\s\S]*?</style>", " ", no_script, flags=re.I)
        return _tidy(re.sub(r"<[^>]+>", " ", no_style))
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return _tidy(soup.get_text("\n", strip=True))


def _tidy(text: str) -> str:
    """Collapse extraction artifacts into stable, chunkable text.

    Deterministic and idempotent: ``_tidy(_tidy(x)) == _tidy(x)``, so the same
    bytes always chunk identically across runs and platforms.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" ?\n ?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Providers
# ---------------------------------------------------------------------------


_FORMAT_BY_EXTENSION: dict[str, str] = {
    **dict.fromkeys(PDF_EXTENSIONS, "pdf"),
    **dict.fromkeys(DOCX_EXTENSIONS, "docx"),
    **dict.fromkeys(PPTX_EXTENSIONS, "pptx"),
    **dict.fromkeys(XLSX_EXTENSIONS, "xlsx"),
    **dict.fromkeys(DOC_EXTENSIONS, "doc"),
    **dict.fromkeys(RTF_EXTENSIONS, "rtf"),
    **dict.fromkeys(EPUB_EXTENSIONS, "epub"),
    **dict.fromkeys(HTML_EXTENSIONS, "html"),
}


def _format_for(relative_path: str) -> str | None:
    return _FORMAT_BY_EXTENSION.get(Path(relative_path).suffix.lower())


def extract_builtin(kind: str | None, content: bytes) -> str:
    """Standard-library-only extraction for one already-resolved format.

    The single place formats map to readers, so a new format cannot be wired
    into one provider and silently forgotten in another. PPTX/XLSX/EPUB/RTF
    live in :mod:`pheasant.ingestion.office` and legacy DOC in
    :mod:`pheasant.ingestion.msdoc`; imports are local so a text-only region
    never loads any of them.
    """
    if kind == "pdf":
        return extract_pdf_text_builtin(content)
    if kind == "docx":
        return extract_docx_text_builtin(content)
    if kind == "html":
        return extract_html_text(content)
    if kind == "pptx":
        from pheasant.ingestion.office import extract_pptx_text

        return _tidy(extract_pptx_text(content))
    if kind == "xlsx":
        from pheasant.ingestion.office import extract_xlsx_text

        return _tidy(extract_xlsx_text(content))
    if kind == "epub":
        from pheasant.ingestion.office import extract_epub_text

        return _tidy(extract_epub_text(content))
    if kind == "rtf":
        from pheasant.ingestion.office import extract_rtf_text

        return _tidy(extract_rtf_text(content))
    if kind == "doc":
        from pheasant.ingestion.msdoc import extract_doc_text

        return _tidy(extract_doc_text(content))
    return ""


class BuiltinExtractor:
    """Standard-library-only extractor. No third-party imports, ever."""

    def __init__(self, model: str = "builtin"):
        self.model = model
        self.calls = 0

    def extract(self, content: bytes, relative_path: str, sidecar: bytes | None = None) -> str:
        authored = _sidecar_text(sidecar)
        if authored is not None:
            return authored
        self.calls += 1
        return extract_builtin(_format_for(relative_path), content)


class NativeExtractor:
    """``pymupdf`` + ``python-docx`` extractor (both already core deps).

    Higher fidelity than the builtin on PDFs with CID/Type0 fonts, custom
    encodings, or complex layout. Each format's import is deferred to call
    time and falls back to the builtin, so a missing optional wheel degrades
    that one format instead of failing the sync.
    """

    def __init__(self, model: str = "native"):
        self.model = model
        self.calls = 0
        self._builtin = BuiltinExtractor()

    def extract(self, content: bytes, relative_path: str, sidecar: bytes | None = None) -> str:
        authored = _sidecar_text(sidecar)
        if authored is not None:
            return authored
        self.calls += 1
        kind = _format_for(relative_path)
        if kind == "pdf":
            return self._extract_pdf(content, relative_path)
        if kind == "docx":
            return self._extract_docx(content, relative_path)
        if kind == "epub":
            return self._extract_epub(content, relative_path)
        # PPTX, XLSX, RTF and legacy DOC have no third-party reader in this
        # project's dependency tree (`python-docx` is OOXML-word-only,
        # `pymupdf` does not open them), and the builtin readers are complete
        # for them — the OOXML/EPUB text *is* the XML, and RTF is a text
        # format. So "native" is honestly the same code path here rather than
        # a pretend upgrade.
        return extract_builtin(kind, content)

    def _extract_pdf(self, content: bytes, relative_path: str) -> str:
        try:
            import pymupdf  # type: ignore[import-not-found]
        except ModuleNotFoundError:
            try:
                import fitz as pymupdf  # type: ignore[import-not-found,no-redef]
            except ModuleNotFoundError:
                logger.debug("pymupdf unavailable; using builtin PDF extraction")
                return extract_pdf_text_builtin(content)
        try:
            with pymupdf.open(stream=content, filetype="pdf") as document:
                pages = [page.get_text() or "" for page in document]
            return _tidy("\n\n".join(pages))
        except Exception as exc:
            # A corrupt/encrypted PDF must not abort a sync: fall back to the
            # builtin, which may still recover the uncompressed streams.
            logger.warning("pymupdf failed on %s (%s); using builtin", relative_path, exc)
            return extract_pdf_text_builtin(content)

    def _extract_epub(self, content: bytes, relative_path: str) -> str:
        """EPUB via pymupdf, which reads it natively (not just PDF).

        A genuine upgrade over the builtin reader rather than a rename: MuPDF
        lays the book out and walks it in reading order, so it recovers text
        from packages whose spine or manifest the builtin path cannot follow.
        Falls back to the builtin reader on any failure.
        """
        try:
            import pymupdf  # type: ignore[import-not-found]
        except ModuleNotFoundError:
            try:
                import fitz as pymupdf  # type: ignore[import-not-found,no-redef]
            except ModuleNotFoundError:
                return extract_builtin("epub", content)
        try:
            with pymupdf.open(stream=content, filetype="epub") as document:
                pages = [page.get_text() or "" for page in document]
            text = _tidy("\n\n".join(pages))
            return text or extract_builtin("epub", content)
        except Exception as exc:
            logger.warning("pymupdf failed on %s (%s); using builtin", relative_path, exc)
            return extract_builtin("epub", content)

    def _extract_docx(self, content: bytes, relative_path: str) -> str:
        import io

        try:
            import docx  # type: ignore[import-not-found]
        except ModuleNotFoundError:
            logger.debug("python-docx unavailable; using builtin DOCX extraction")
            return extract_docx_text_builtin(content)
        try:
            document = docx.Document(io.BytesIO(content))
            lines = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    lines.extend(cell.text for cell in row.cells)
            return _tidy("\n".join(lines))
        except Exception as exc:
            logger.warning("python-docx failed on %s (%s); using builtin", relative_path, exc)
            return extract_docx_text_builtin(content)


class AutoExtractor:
    """Default provider: native when it yields text, builtin otherwise.

    "Yields text" rather than "imports successfully" is deliberate — a PDF
    whose fonts pymupdf renders as empty may still have readable
    uncompressed streams, and vice versa. Trying both and keeping the longer
    result costs one extra pass on the rare empty case and is the difference
    between indexing a document and silently dropping it.
    """

    def __init__(self, model: str = "auto"):
        self.model = model
        self.calls = 0
        self._native = NativeExtractor()
        self._builtin = BuiltinExtractor()

    def extract(self, content: bytes, relative_path: str, sidecar: bytes | None = None) -> str:
        authored = _sidecar_text(sidecar)
        if authored is not None:
            return authored
        self.calls += 1
        try:
            text = self._native.extract(content, relative_path)
        except Exception as exc:
            logger.warning("native extraction failed on %s: %s", relative_path, exc)
            text = ""
        if text.strip():
            return text
        try:
            return self._builtin.extract(content, relative_path)
        except Exception as exc:
            logger.warning("builtin extraction failed on %s: %s", relative_path, exc)
            return ""


def build_extractor(settings: ExtractorSettings) -> DocumentExtractor:
    provider = (settings.provider or "auto").lower()
    if provider == "auto":
        return AutoExtractor()
    if provider == "builtin":
        return BuiltinExtractor()
    if provider == "native":
        return NativeExtractor()
    if provider == "sandboxed":
        from pheasant.ingestion.extractor_sandbox import SandboxedExtractor

        return SandboxedExtractor(fallback=AutoExtractor())
    raise ValueError(
        f"Unsupported ingestion.extractor.provider {settings.provider!r}; "
        "expected 'auto', 'native', 'builtin' or 'sandboxed'"
    )


def extractable_extensions(settings: ExtractorSettings | None = None) -> set[str]:
    """Extensions this extractor handles for a given config.

    HTML is opt-in: ``.html``/``.xml`` are long-standing members of
    ``TEXT_EXTENSIONS`` and already index as raw markup, so turning stripping
    on changes the indexed text (and therefore chunk boundaries) of existing
    deployments. Gating it behind ``html_text`` keeps the default byte-identical
    and makes the improvement an explicit operator choice.
    """
    extensions = set(EXTRACTED_EXTENSIONS)
    if settings is not None and getattr(settings, "html_text", False):
        extensions |= set(HTML_EXTENSIONS)
    return extensions


def source_includes_documents(source: Any) -> bool:
    """True when a source's ``include`` globs admit any extractable document.

    Document extraction is opt-in exactly like images/audio: the default
    ``include`` list is code/markdown/config only, so a region grows document
    capability only when an operator adds e.g. ``**/*.pdf`` or ``**/*.pptx``.

    This must cover **every** extension in ``EXTRACTED_EXTENSIONS``. A format
    the pipeline accepts but this predicate misses would build no extractor for
    a source carrying only that format, so its files would index with zero
    chunks — the original bug, reintroduced one format at a time. A test
    asserts the two sets agree.
    """
    includes = list(getattr(source, "include", None) or [])
    admitted = tuple(EXTRACTED_EXTENSIONS)
    return any(pattern.lower().endswith(admitted) for pattern in includes)


def config_has_document_source(config: Any) -> bool:
    return any(source_includes_documents(source) for source in getattr(config, "sources", []))


def extractor_from_config(config: Any, *, source: Any = None) -> DocumentExtractor | None:
    """Build the document extractor, or ``None`` when no document source exists.

    Returning ``None`` keeps a text-only region byte-identical to its
    pre-extraction behavior: no extractor is constructed and
    ``read_text``/``read_text_bytes`` behave exactly as before.

    Pass ``source`` to ask about one source rather than the whole config. That
    is what a source registered *after* the engine was built needs — the config
    the engine was constructed from never mentioned it.
    """
    settings = getattr(getattr(config, "ingestion", None), "extractor", None)
    if settings is None:
        return None
    html_on = bool(getattr(settings, "html_text", False))
    wanted = (
        source_includes_documents(source) if source is not None else config_has_document_source(config)
    )
    if not wanted and not html_on:
        return None
    try:
        return build_extractor(settings)
    except ValueError as exc:
        logger.warning("Document extraction disabled: %s", exc)
        return None
